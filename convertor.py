import os
from docx2pdf import convert
from pdf2docx import Converter
import docx
from fpdf import FPDF
import fitz
import pandas as pd
from PIL import Image

#------------------------------------------------------------------------------------

def word_to_pdf(input_file, output_file):
    # Check if the input file exists
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        return

    # Convert the Word document to PDF
    convert(input_file, output_file)
    print(f"File converted successfully: {output_file}")

#------------------------------------------------------------------------------------

def pdf_to_word(input_file, output_file):
    # Check if the input file exists
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        return

    # Convert PDF to Word
    cv = Converter(input_file)
    cv.convert(output_file, start=0, end=None)
    cv.close()
    print(f"File converted successfully: {output_file}")

#------------------------------------------------------------------------------------

def text_to_word(input_file, output_file):

    # Check if the input file exists
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        return

    with open(input_file, 'r') as file:
        text = file.read()

    doc = docx.Document()
    doc.add_paragraph(text)
    doc.save(output_file)
    print(f"File converted successfully: {output_file}")

#------------------------------------------------------------------------------------

def text_to_pdf(input_file, output_file):

    # Check if the input file exists
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        return
    
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    with open(input_file, 'r') as file:
        for line in file:
            pdf.cell(200, 10, txt=line, ln=True, align='L')

    pdf.output(output_file)
    print(f"File converted successfully: {output_file}")

#------------------------------------------------------------------------------------

def word_to_text(input_file, output_file):

    # Check if the input file exists
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        return
    
    doc = docx.Document(input_file)
    with open(output_file, 'w') as file:
        for paragraph in doc.paragraphs:
            file.write(paragraph.text + '\n')
    print(f"File converted successfully: {output_file}")

#------------------------------------------------------------------------------------

def pdf_to_text(input_file, output_file):

    # Check if the input file exists
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        return
    
    pdf_document = fitz.open(input_file)
    with open(output_file, 'w') as file:
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            file.write(page.get_text())
    print(f"File converted successfully: {output_file}")

#------------------------------------------------------------------------------------

def text_to_json(input_file, output_file):

    # Check if the input file exists
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        return
    
    with open(input_file, 'r') as file:
        lines = file.readlines()

    # Assuming each line in the text file is a separate record
    df = pd.DataFrame({'text': lines})

    # Write to a JSON file
    df.to_json(output_file, orient='records', lines=True)
    print(f"File converted successfully: {output_file}")

#------------------------------------------------------------------------------------

def image_to_pdf(input_file, output_file):

    # Check if the input file exists
    if not os.path.exists(input_file):
        print(f"File not found: {input_file}")
        return
    
    # Open the image file
    image = Image.open(input_file)

    # Ensure the image is in RGB mode (required for saving as PDF)
    if image.mode != 'RGB':
        image = image.convert('RGB')

    # Save the image as a PDF
    image.save(output_file, "PDF", resolution=100.0)
    print(f"Image converted to PDF and saved as: {output_file}")

#------------------------------------------------------------------------------------

def convertor():
    print("word_to_pdf press 1 ")
    print("pdf_to_word press 2 ")
    print("text_to_word press 3 ")
    print("text_to_pdf press 4 ")
    print("word_to_text press 5 ")
    print("pdf_to_text press 6 ")
    print("text_to_json press 7 ")
    print("image_to_pdf press 8 ")
    print("Exit press 0")
    print("________________________________")
    val=int(input("press your choice:  "))
    if (val==1):
        word_to_pdf(input_path, output_path)
    elif(val==2):
        pdf_to_word(input_path, output_path)
    elif(val==3):
        text_to_word(input_path, output_path)
    elif(val==4):
        text_to_pdf(input_path, output_path)
    elif(val==5):
        word_to_text(input_path, output_path)
    elif(val==6):
        pdf_to_text(input_path, output_path)
    elif(val==7):
        text_to_json(input_path, output_path)
    elif(val==8):
        image_to_pdf(input_path, output_path)
    elif(val==0):
        print("Exit")
    else:
        print("your choice is invalid")
        convertor()

#------------------------------------------------------------------------------------

# Take file locations from user input
input_path = input("Enter the input file path (/) format: ")
output_path = input("Enter the output file path (/) format: ")

print(f"Input file path: {input_path}")
print(f"Output file path: {output_path}")
print("--------------------------------------------")

convertor()
#------------------------------------------------------------------------------------